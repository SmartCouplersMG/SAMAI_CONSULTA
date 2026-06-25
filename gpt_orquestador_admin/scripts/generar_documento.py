"""
generar_documento.py
---------------------
Genera la providencia en Word (.docx) y, opcionalmente, en PDF.
POR AHORA SIN MEMBRETE: formato limpio (Times New Roman 12, márgenes y numeración de página).

El membrete (logo, encabezado/pie institucional, sello) se podrá añadir luego en
`configurar_documento` y `crear_word_providencia` usando una plantilla en ../templates.

Uso rápido:
    python generar_documento.py
"""

from __future__ import annotations

import subprocess
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


def configurar_documento(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3)
    section.right_margin = Cm(3)

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)

    for name, size in [("Heading 1", 14), ("Heading 2", 13), ("Heading 3", 12)]:
        style = doc.styles[name]
        style.font.name = "Times New Roman"
        style.font.bold = True
        style.font.size = Pt(size)


def agregar_numero_pagina(section) -> None:
    paragraph = section.footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("Página ")

    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")

    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)


def agregar_parrafo_justificado(doc: Document, texto: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(texto)
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)
    return p


def _es_titulo(linea: str) -> bool:
    encabezados = {"antecedentes", "consideraciones", "resuelve", "falla", "decisión", "decision"}
    return (
        len(linea) < 120
        and (
            linea.isupper()
            or linea.startswith(
                ("I.", "II.", "III.", "IV.", "V.", "VI.", "VII.", "VIII.", "IX.", "X.")
            )
            or linea.lower() in encabezados
        )
    )


def agregar_bloque_texto(doc: Document, texto: str) -> None:
    """Convierte una providencia en texto plano a Word, detectando encabezados simples."""
    for raw in texto.splitlines():
        linea = raw.strip()
        if not linea:
            continue
        if _es_titulo(linea):
            doc.add_heading(linea, level=1)
        else:
            agregar_parrafo_justificado(doc, linea)


def crear_word_providencia(
    texto_providencia: str,
    salida_docx: str = "output/providencia.docx",
    titulo: str = "PROYECTO DE PROVIDENCIA",
) -> str:
    Path(salida_docx).parent.mkdir(parents=True, exist_ok=True)

    doc = Document()
    configurar_documento(doc)
    agregar_numero_pagina(doc.sections[0])

    # --- Aquí, en el futuro, va el MEMBRETE (logo + encabezado institucional) ---

    titulo_p = doc.add_paragraph()
    titulo_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = titulo_p.add_run(titulo)
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(14)

    doc.add_paragraph()
    agregar_bloque_texto(doc, texto_providencia)

    doc.save(salida_docx)
    return salida_docx


def convertir_a_pdf(docx_path: str) -> str:
    """Convierte a PDF con LibreOffice (debe estar instalado y en el PATH)."""
    docx_path = Path(docx_path)
    out_dir = docx_path.parent
    subprocess.run(
        ["soffice", "--headless", "--convert-to", "pdf", "--outdir", str(out_dir), str(docx_path)],
        check=True,
    )
    return str(out_dir / f"{docx_path.stem}.pdf")


PROVIDENCIA_DEMO = """
JUZGADO ADMINISTRATIVO ORAL DEL CIRCUITO DE [CIUDAD]

PROCESO: REPARACIÓN DIRECTA
DEMANDANTE: [NOMBRE]
DEMANDADO: [ENTIDAD]
RADICADO: [RADICADO]

I. ANTECEDENTES

La parte demandante solicita que se declare patrimonialmente responsable a la entidad demandada
por los daños antijurídicos presuntamente causados por acción u omisión de sus agentes.

II. CONSIDERACIONES

El despacho analizará la procedencia de la demanda, la existencia del daño antijurídico, la
imputación jurídica y fáctica, el nexo causal y la prueba de los perjuicios reclamados.

III. RESUELVE

PRIMERO. [Decisión principal].
SEGUNDO. [Decisión sobre costas].
TERCERO. Ejecutoriada esta providencia, archívese el expediente.
"""


if __name__ == "__main__":
    docx = crear_word_providencia(PROVIDENCIA_DEMO)
    print(f"Documento Word creado: {docx}")
    # Descomenta si tienes LibreOffice instalado:
    # pdf = convertir_a_pdf(docx)
    # print(f"PDF creado: {pdf}")
